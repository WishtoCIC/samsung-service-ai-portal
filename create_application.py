from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 페이지 설정 ───────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)

FONT_NAME = '맑은 고딕'

# ── 헬퍼 함수 ─────────────────────────────────────────
def font(run, size=10, bold=False, color=None, italic=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)

def para(container, text='', size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
         sb=0, sa=4, color=None, italic=False):
    if hasattr(container, 'add_paragraph'):
        p = container.add_paragraph()
    else:
        p = container
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        r = p.add_run(text)
        font(r, size=size, bold=bold, color=color, italic=italic)
    return p

def cell_para(cell, text='', size=10, bold=False,
              align=WD_ALIGN_PARAGRAPH.LEFT, color=None, italic=False,
              sb=2, sa=2, clear=True):
    if clear:
        cell.text = ''
    p = cell.paragraphs[0] if clear else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if text:
        r = p.add_run(text)
        font(r, size=size, bold=bold, color=color, italic=italic)
    return p

def add_run_to_cell(cell, text, size=10, bold=False, color=None, italic=False, first=False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    if not first:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    font(r, size=size, bold=bold, color=color, italic=italic)
    return p

def shade_cell(cell, hex_color='D9D9D9'):
    tc  = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

def set_tblGrid(table, col_widths_cm):
    tbl = table._tbl
    existing = tbl.find(qn('w:tblGrid'))
    if existing is not None:
        tbl.remove(existing)
    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths_cm:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w * 567)))
        tblGrid.append(gridCol)
    tbl.insert(1, tblGrid)

def set_table_fixed(table, total_cm=15.0):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # 전체 너비 고정
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(int(total_cm * 567)))
    tblW.set(qn('w:type'), 'dxa')
    ex = tblPr.find(qn('w:tblW'))
    if ex is not None: tblPr.remove(ex)
    tblPr.append(tblW)
    # 레이아웃 고정
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    ex2 = tblPr.find(qn('w:tblLayout'))
    if ex2 is not None: tblPr.remove(ex2)
    tblPr.append(tblLayout)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        cell = row.cells[col_idx]
        cell.width = Cm(width_cm)
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(int(width_cm * 567)))
        tcW.set(qn('w:type'), 'dxa')
        ex = tcPr.find(qn('w:tcW'))
        if ex is not None: tcPr.remove(ex)
        tcPr.append(tcW)

def min_row_height(row, cm):
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'),   str(int(cm * 567)))
    trH.set(qn('w:hRule'), 'atLeast')
    existing = trPr.find(qn('w:trHeight'))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trH)

def set_table_border(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tblBorders.append(el)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

# ════════════════════════════════════════════════════════
#  제목
# ════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(18)
r = p_title.add_run('사내공모 지원서')
font(r, size=20, bold=True)

# ════════════════════════════════════════════════════════
#  1. 지원자 인적사항
# ════════════════════════════════════════════════════════
para(doc, '1. 지원자 인적사항', size=11, bold=True, sb=0, sa=6)

# 2행 4열 테이블
t1 = doc.add_table(rows=2, cols=4)
set_table_border(t1)
set_table_fixed(t1, 15.0)
set_tblGrid(t1, [1.8, 5.7, 1.8, 5.7])

labels1 = [['소속', '운영혁신팀 구독혁신그룹', '사번', ''],
           ['이름', '',                        '현 직무', '구독 상품화 담당 및 판매연계 활성화 업무지원']]

for r_i, row_data in enumerate(labels1):
    row = t1.rows[r_i]
    for c_i, text in enumerate(row_data):
        cell = row.cells[c_i]
        is_label = (c_i % 2 == 0)
        shade_cell(cell, 'D9D9D9') if is_label else None
        cell_para(cell, text, size=10, bold=is_label,
                  align=WD_ALIGN_PARAGRAPH.CENTER, sb=4, sa=4)

# 열 너비: 레이블 좁게, 값 넓게 (전체 15cm = A4 - 좌우여백)
for col_i, w in enumerate([1.8, 5.7, 1.8, 5.7]):
    set_col_width(t1, col_i, w)

# ════════════════════════════════════════════════════════
#  2. 지원분야
# ════════════════════════════════════════════════════════
para(doc, '2. 지원분야', size=11, bold=True, sb=14, sa=6)

t2 = doc.add_table(rows=2, cols=2)
set_table_border(t2)
set_table_fixed(t2, 15.0)
set_tblGrid(t2, [3.0, 12.0])

# 헤더행
shade_cell(t2.rows[0].cells[0], 'D9D9D9')
shade_cell(t2.rows[0].cells[1], 'D9D9D9')
cell_para(t2.rows[0].cells[0], '구분',  size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sb=3, sa=3)
cell_para(t2.rows[0].cells[1], '내용',  size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sb=3, sa=3)

# 데이터행
shade_cell(t2.rows[1].cells[0], 'EEF2F7')
cell_para(t2.rows[1].cells[0], '지원분야', size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sb=4, sa=4)
cell_para(t2.rows[1].cells[1], '프로세스 혁신 (모바일/디지털)', size=10, bold=True, sb=4, sa=4)

set_col_width(t2, 0, 3.0)
set_col_width(t2, 1, 12.0)

# 안내 문구
p_note2 = para(doc, sb=3, sa=14)
r1 = p_note2.add_run('※ 지원분야 : ')
font(r1, size=9, color=(80,80,80))
r2 = p_note2.add_run('프로세스 혁신(모바일/디지털), 시스템, 교육/문화, AI Crew 운영 중 택일')
font(r2, size=9, color=(80,80,80))

# ════════════════════════════════════════════════════════
#  3. 지원 내용
# ════════════════════════════════════════════════════════
para(doc, '3. 지원 내용', size=11, bold=True, sb=0, sa=6)

t3 = doc.add_table(rows=4, cols=2)
set_table_border(t3)
set_table_fixed(t3, 15.0)
set_tblGrid(t3, [3.0, 12.0])

# 헤더행
shade_cell(t3.rows[0].cells[0], 'D9D9D9')
shade_cell(t3.rows[0].cells[1], 'D9D9D9')
cell_para(t3.rows[0].cells[0], '구분', size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sb=3, sa=3)
cell_para(t3.rows[0].cells[1], '내용', size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sb=3, sa=3)

# ── 행1: 지원동기 ──────────────────────────────────────
shade_cell(t3.rows[1].cells[0], 'EEF2F7')
c0 = t3.rows[1].cells[0]
c0.text = ''
p_ = c0.paragraphs[0]
p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_.paragraph_format.space_before = Pt(4)
p_.paragraph_format.space_after  = Pt(2)
c0.vertical_alignment = WD_ALIGN_VERTICAL.TOP
r_ = p_.add_run('지원동기')
font(r_, size=10, bold=True)
p2_ = c0.add_paragraph()
p2_.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2_.paragraph_format.space_before = Pt(0)
p2_.paragraph_format.space_after  = Pt(4)
r2_ = p2_.add_run('* 상세히 작성')
font(r2_, size=8, color=(120,120,120), italic=True)

c1 = t3.rows[1].cells[1]
c1.text = ''
c1.vertical_alignment = WD_ALIGN_VERTICAL.TOP

motiv_texts = [
    "현장 엔지니어로 고객을 직접 만나던 시절부터 현재 구독 업무에 이르기까지, 직무는 달라졌지만 반복적으로 마주친 문제가 있었습니다. 기술 가이드, 수리 매뉴얼, 사양 변경 이력 등 방대한 자료가 존재함에도 여러 곳에 분산되어 있어 필요한 순간에 신속하게 찾기 어려웠고, 이로 인해 현장 대응이 지연되거나 업무 흐름이 끊기는 상황이 반복되었습니다. 정보의 부재가 아니라 프로세스의 문제였습니다.",
    "이를 해결할 수 있다는 가능성을 육아휴직 기간 동안 직접 확인했습니다. Claude·ChatGPT를 활용한 바이브 코딩을 시작으로, LangChain·LangGraph 기반의 RAG(검색 증강 생성) 시스템을 설계·구축하며 분산된 정보를 AI가 즉시 찾아주는 구조를 실습으로 검증했습니다. 아울러 사내 보안 환경을 고려하여 외부 LLM 의존도를 낮춘 로컬 LLM(Gemma4) 온프레미스 환경도 직접 테스트하며 현실적인 사내 적용 방안을 검토했습니다.",
    "지원 과정에서도 이 역량을 직접 시연해 보았습니다. 기상청 공공데이터(기온·폭염일수)와 에어컨 서비스 물량 이력을 연계한 수요예측 대시보드를 Python·Streamlit으로 직접 구현하였습니다. 분석 결과 7~8월 폭염일수와 서비스 접수량 간 상관계수 r=0.66이 확인되었고, 제품코드(HAC/SRA/CAC)별 월별 트렌드·지역 GPS 분포·처리시간 분석까지 실제로 작동하는 프로토타입을 완성하였습니다. 말로만 하는 혁신이 아니라, 이미 만들어봤습니다.",
    "현장·교육·기술·운영을 순서대로 경험하며 쌓아온 현업 시각과, 스스로 준비한 AI 역량을 결합하여 실제로 작동하는 프로세스 혁신에 기여하고 싶어 지원합니다.",
]
for i, text in enumerate(motiv_texts):
    p = c1.paragraphs[0] if i == 0 else c1.add_paragraph()
    p.paragraph_format.space_before = Pt(4 if i == 0 else 0)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    font(r, size=10)

min_row_height(t3.rows[1], 5.0)

# ── 행2: 직무경험/자격/수상사항 ───────────────────────
shade_cell(t3.rows[2].cells[0], 'EEF2F7')
c0b = t3.rows[2].cells[0]
c0b.text = ''
c0b.vertical_alignment = WD_ALIGN_VERTICAL.TOP
pb = c0b.paragraphs[0]
pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
pb.paragraph_format.space_before = Pt(4)
pb.paragraph_format.space_after  = Pt(2)
rb = pb.add_run('직무경험 /\n자격 /\n수상사항')
font(rb, size=10, bold=True)

c1b = t3.rows[2].cells[1]
c1b.text = ''
c1b.vertical_alignment = WD_ALIGN_VERTICAL.TOP

def add_p(cell, text, size=10, bold=False, sb=0, sa=3, indent=False):
    p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
    # 만약 첫 문단이 비어있으면 재활용, 아니면 추가
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].runs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.left_indent = Pt(10)
    r = p.add_run(text)
    font(r, size=size, bold=bold)
    return p

# 직무경험 내용 직접 작성
c1b.text = ''
c1b.vertical_alignment = WD_ALIGN_VERTICAL.TOP

exp_entries = [
    ('[ 직무 경험 ]', True, False, 4, 4),
    ('• 운영혁신팀 구독혁신그룹 (현재)   |   구독 상품화 및 판매연계 활성화 업무지원', False, True, 0, 3),
    ('• 기술팀 디지털기술그룹   |   세탁기 제품 담당, 기술 자료 관리 및 제품 지원', False, True, 0, 3),
    ('• 기술팀 기술교육그룹   |   HA 제품군 기술강사, 교육 콘텐츠 제작·기술교육 운영', False, True, 0, 3),
    ('• 현장 엔지니어   |   고객 방문 제품 수리, 현장 프로세스 직접 경험', False, True, 0, 6),
    ('[ AI 자기개발 — 육아휴직 중 ]', True, False, 0, 4),
    ('• LangChain·LangGraph 기반 RAG 시스템 설계·구축 실습', False, True, 0, 3),
    ('• 로컬 LLM(Gemma4) 온프레미스 환경 테스트 — 사내 보안 제약 대응 방안 자체 검토', False, True, 0, 3),
    ('• Claude·ChatGPT 활용 바이브 코딩 학습', False, True, 0, 4),
]

for i, (text, bold_, indent_, sb_, sa_) in enumerate(exp_entries):
    if i == 0:
        p = c1b.paragraphs[0]
        p.paragraph_format.space_before = Pt(sb_)
        p.paragraph_format.space_after  = Pt(sa_)
    else:
        p = c1b.add_paragraph()
        p.paragraph_format.space_before = Pt(sb_)
        p.paragraph_format.space_after  = Pt(sa_)
    if indent_:
        p.paragraph_format.left_indent = Pt(8)
    r = p.add_run(text)
    font(r, size=10, bold=bold_)

min_row_height(t3.rows[2], 5.0)

# ── 행3: 기타 ─────────────────────────────────────────
shade_cell(t3.rows[3].cells[0], 'EEF2F7')
cell_para(t3.rows[3].cells[0], '기타', size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sb=4, sa=4)
t3.rows[3].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

c1c = t3.rows[3].cells[1]
c1c.text = ''
c1c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
p_etc = c1c.paragraphs[0]
p_etc.paragraph_format.space_before = Pt(4)
p_etc.paragraph_format.space_after  = Pt(4)
etc_entries = [
    ('[ 프로토타입 ① — 서비스 수요예측 대시보드 ]', True, 4, 4),
    ('기상청 공공 API(기온·폭염일수) × 에어컨 서비스 이력 데이터를 연계하여 월별 수요를 '
     '예측하는 대시보드를 Python·Streamlit으로 직접 구현하였습니다. 제품코드별(HAC·SRA·CAC) '
     '월별 트렌드, 지역 GPS 분포 지도, 날씨 상관관계(r=0.66), 처리시간 분포 등 5개 탭을 '
     '실제로 작동하는 형태로 완성하였으며, 현장에서 실 데이터를 CSV/XLSX로 업로드하면 '
     '즉시 분석 결과로 전환되는 구조로 설계하였습니다.', False, 0, 8),
    ('[ 프로토타입 ② — 기술정보 통합 검색 웹 앱 (구상·구현 중) ]', True, 0, 4),
    ('수리 매뉴얼·제품 사양·사양 변경 이력 등 흩어진 기술 자료를 RAG(검색 증강 생성) 방식으로 '
     '통합 검색하고, 현장 엔지니어가 담당자에게 직접 요청을 접수할 수 있는 웹 기반 앱을 '
     'LangChain·Streamlit으로 설계하고 있습니다. 사내 보안 제약을 고려하여 로컬 LLM(Gemma4) '
     '기반 온프레미스 구성을 우선 검토 중이며, TF 합류 시 즉시 사내 환경에 적용 가능한 '
     '형태로 발전시킬 계획입니다.', False, 0, 8),
    ('현장·교육·기술·운영 전 단계를 직접 경험한 시각으로, 실제로 작동하는 AI 기반 프로세스 '
     '혁신 과제 발굴과 구현에 즉시 기여하겠습니다.', False, 0, 4),
]

for i, entry in enumerate(etc_entries):
    if len(entry) == 4:
        text, bold_, sb_, sa_ = entry
    else:
        text, bold_, sb_, sa_ = entry[0], False, 0, 4

    if i == 0:
        p_etc2 = c1c.paragraphs[0]
    else:
        p_etc2 = c1c.add_paragraph()
    p_etc2.paragraph_format.space_before = Pt(sb_)
    p_etc2.paragraph_format.space_after  = Pt(sa_)
    r_etc2 = p_etc2.add_run(text)
    font(r_etc2, size=10, bold=bold_)

min_row_height(t3.rows[3], 2.5)

# 열 너비
set_col_width(t3, 0, 3.0)
set_col_width(t3, 1, 12.0)

# 안내 문구
p_foot = para(doc, sb=6, sa=0)
r_foot = p_foot.add_run('* 내용이 많을 경우 페이지를 추가하여 작성 가능합니다.')
font(r_foot, size=9, color=(120,120,120), italic=True)

# ── 저장 ──────────────────────────────────────────────
output = '/Users/user/00_AI_WORKS/AI_TF지원/AI_TF_사내공모_지원서.docx'
doc.save(output)
print(f'완료: {output}')
